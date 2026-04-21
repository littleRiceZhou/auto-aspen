import docx
import os
from pathlib import Path
from docx.shared import RGBColor, Inches, Pt
from docx.oxml.shared import qn
import re

file_dir = "static/re"
os.makedirs(file_dir, exist_ok=True)

# 特殊变量字体大小配置
special_font_size = {
    "auto_aspen_26": 16,
    "auto_aspen_8": 8,
    # "auto_aspen_time": 14,
    # "auto_aspen_7": 12,
    # "auto_aspen_5": 18,
}

def get_special_font_config():
    """
    获取特殊变量的字体大小配置
    
    Returns:
        dict: 包含特殊变量配置的字典
    """
    config = {
        "special_font_size": special_font_size
    }
    
    print(f"🔧 特殊字体配置: {config}")
    return config

def get_special_font_for_variable(variable_name):
    """
    获取特定变量的特殊字体大小
    
    Args:
        variable_name (str): 变量名（如 'auto_aspen_26'）
    
    Returns:
        float or None: 特殊字体大小，如果不是特殊变量则返回None
    """
    config = get_special_font_config()
    return config.get('special_font_size', {}).get(variable_name)

def get_auto_aspen_parameter_mapping():
    """
    根据文档中的auto_aspen参数创建完整的映射字典
    
    Returns:
        dict: 参数映射字典，键为auto_aspen参数名，值为对应的实际数值
    """
    return {
        # 接入参数 - 天然气处理机组
        "auto_aspen_1": "50000",      # 最大气量 (m³/d)
        "auto_aspen_2": "13.5",       # 进站压力 (MPaA)
        "auto_aspen_3": "25",         # 平均进气温度 (℃)
        "auto_aspen_4": "6.0",        # 出站压力 (MPaA)
        "auto_aspen_27": "135",       # 进站压力 (MPaG)，与 MPaA×10 同量级，供模板占位
        "auto_aspen_28": "60",       # 出站压力 (MPaG)
        
        # 段落中的参数
        "auto_aspen_5": "654",        # 净发电功率 (kW)
        
        # 机组参数表
        "auto_aspen_6": "1",          # 透平机头数
        "auto_aspen_7": "TRT-1000",   # 机组型号
        "auto_aspen_8": "1",          # 级数
        "auto_aspen_9": "45",         # 机组排气温度 (℃)
        
        # 机组占地面积与操作重量
        "auto_aspen_10": "TRT-1000-EX", # 机组型号
        "auto_aspen_11": "8.5×3.2×3.8", # 机组整体外形尺寸 (长×宽×高 m)
        "auto_aspen_12": "12000",     # 机组整体重量/整体维修保养最大重量 (kg)
        
        # 项目整体经济效益核算
        "auto_aspen_13": "525.6",     # 年净发电量 (kWh)
        "auto_aspen_14": "315.36",    # 年净发电收益 (万元)
        "auto_aspen_15": "184.0",     # 年节约标准煤 (吨)
        "auto_aspen_16": "505.4",     # 年减少CO₂排放 (吨)
        
        # 机组公用工程 - 电源设备参数
        "auto_aspen_17": "15",        # 辅油泵功率 (kW)
        "auto_aspen_18": "12",        # 润滑油电加热器功率 (kW)
        "auto_aspen_19": "8",         # 油雾分离器功率 (kW)
        "auto_aspen_20": "1",         # 发电机加热器=空间加热器功率 (kW)
        "auto_aspen_21": "3",         # PLC柜功率 (kW)
        
        # 机组公用工程 - 水油气参数
        "auto_aspen_22": "850",       # 油冷器流量 (m³/Hr)
        "auto_aspen_23": "45",        # 齿轮箱润滑油流量 (L/min)
        "auto_aspen_24": "120",       # 氮气-干气密封气体流量 (Nm³/h)
        "auto_aspen_25": "95",        # 压缩空气-气动阀气体流量 (Nm³/h)
        
        # 用户信息
        "auto_aspen_26": "用户姓名",     # 用户名称
        "auto_aspen_time": "2025-07-05", # 生成时间
    }

def create_replacement_dict(parameter_values=None):
    """
    创建用于文档替换的字典
    
    Args:
        parameter_values (dict, optional): 自定义参数值，如果不提供则使用默认值
    
    Returns:
        dict: 用于替换的字典，键为原文档中的占位符，值为新的数值
    """
    # 获取默认参数映射
    default_params = get_auto_aspen_parameter_mapping()
    
    # 如果提供了自定义值，则更新默认值
    if parameter_values:
        default_params.update(parameter_values)
    
    return default_params

def replace_text_in_paragraph(paragraph, old_text, new_text, force_font_size=None):
    """
    在段落中替换文本，保持原始格式，支持跨run的文本替换
    
    Args:
        paragraph: docx段落对象
        old_text (str): 要替换的文本
        new_text (str): 新文本
        force_font_size (float, optional): 强制设置字体大小（点数），如果不提供则保持原始大小
    
    Returns:
        int: 替换次数
    """
    # 检查段落文本中是否包含要替换的文本
    full_text = paragraph.text
    if old_text not in full_text:
        return 0
    
    # ⚠️ 关键修复：使用精确匹配，避免子字符串污染
    # 检查是否真的存在需要替换的完整匹配
    if old_text.startswith('auto_aspen_'):
        # 确保 auto_aspen_1 不会匹配 auto_aspen_14
        # 模式：auto_aspen_数字，后面必须是非数字字符或字符串结尾
        escaped_text = re.escape(old_text)
        pattern = escaped_text + r'(?=\D|$)'  # 正向前瞻：后面是非数字或字符串结尾
        if not re.search(pattern, full_text):
            print(f"🚫 精确匹配检查：'{old_text}' 在 '{full_text}' 中只是子字符串，跳过替换")
            return 0
    else:
        # 对于其他格式，使用单词边界
        pattern = r'\b' + re.escape(old_text) + r'\b'
        if not re.search(pattern, full_text):
            print(f"🚫 精确匹配检查：'{old_text}' 在 '{full_text}' 中只是子字符串，跳过替换")
            return 0
    
    print(f"🔍 段落文本检查: 找到 '{old_text}' 在段落中（精确匹配）")
    print(f"   段落完整文本: '{full_text}'")
    print(f"   段落runs数量: {len(paragraph.runs)}")
    
    # 如果只有一个run，使用简单方法
    if len(paragraph.runs) == 1:
        run = paragraph.runs[0]
        if old_text in run.text:
            return replace_text_in_single_run(run, old_text, new_text, force_font_size)
        return 0
    
    # 多run情况：需要跨run替换
    return replace_text_across_runs(paragraph, old_text, new_text, force_font_size)

def replace_text_in_single_run(run, old_text, new_text, force_font_size=None):
    """
    在单个run中替换文本并保持格式，支持特殊变量字体大小
    """
    if old_text not in run.text:
        return 0
    
    # 保存原始格式
    original_bold = run.bold
    original_italic = run.italic
    original_underline = run.underline
    original_font_name = run.font.name
    original_font_size = run.font.size
    original_font_color = run.font.color.rgb if run.font.color.rgb else None
    
    print(f"🔍 单run替换: '{old_text}' -> '{new_text}'")
    print(f"   原始字体大小: {original_font_size}")
    
    # 执行替换
    old_count = run.text.count(old_text)
    run.text = run.text.replace(old_text, new_text)
    
    # 恢复/设置格式
    if original_bold is not None:
        run.bold = original_bold
    if original_italic is not None:
        run.italic = original_italic
    if original_underline is not None:
        run.underline = original_underline
    if original_font_name is not None:
        run.font.name = original_font_name
    
    # 智能字体大小设置
    final_font_size = None
    font_size_source = ""
    
    # 检查是否为特殊变量
    special_config = get_special_font_config()
    if old_text in special_config.get('special_font_size', {}):
        final_font_size = special_config['special_font_size'][old_text]
        font_size_source = f"特殊变量字体({final_font_size}pt)"
    else:
        # 不是特殊变量，使用常规逻辑
        if force_font_size is not None:
            final_font_size = force_font_size
            font_size_source = f"强制字体({final_font_size}pt)"
        elif original_font_size is not None:
            # 处理原始字体大小
            if hasattr(original_font_size, 'pt'):
                final_font_size = original_font_size.pt
            else:
                final_font_size = original_font_size
            font_size_source = f"原始字体({final_font_size}pt)"
    
    # 应用字体大小
    if final_font_size is not None:
        run.font.size = Pt(final_font_size)
        print(f"🔧 应用字体大小: {font_size_source}")
    
    if original_font_color is not None:
        run.font.color.rgb = original_font_color
    
    print(f"✅ 单run替换完成: {old_count}次")
    return old_count

def replace_text_across_runs(paragraph, old_text, new_text, force_font_size=None):
    """
    跨run替换文本，这是最复杂的情况
    """
    print(f"🔍 跨run替换: '{old_text}' -> '{new_text}'")
    
    # 收集所有runs的信息
    runs_info = []
    char_position = 0
    
    for i, run in enumerate(paragraph.runs):
        run_info = {
            'index': i,
            'text': run.text,
            'start_pos': char_position,
            'end_pos': char_position + len(run.text),
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color.rgb else None,
            'run_obj': run
        }
        runs_info.append(run_info)
        char_position += len(run.text)
    
    # 查找所有匹配位置，使用精确匹配
    full_text = paragraph.text
    matches = []
    
    if old_text.startswith('auto_aspen_'):
        # 对 auto_aspen_ 格式使用精确匹配
        escaped_text = re.escape(old_text)
        pattern = escaped_text + r'(?=\D|$)'  # 正向前瞻：后面是非数字或字符串结尾
        for match in re.finditer(pattern, full_text):
            matches.append((match.start(), match.end()))
    else:
        # 对其他格式使用普通查找
        start = 0
        while True:
            pos = full_text.find(old_text, start)
            if pos == -1:
                break
            matches.append((pos, pos + len(old_text)))
            start = pos + 1
    
    if not matches:
        return 0
    
    print(f"   找到 {len(matches)} 个匹配位置: {matches}")
    
    # ⚠️ 关键：从后往前替换，避免位置偏移！
    # 如果有多个匹配项，从前往后替换会改变后面匹配项的位置
    # 例如：文本"auto_aspen_1和auto_aspen_2"，如果先替换auto_aspen_1为"50000"
    # 会变成"50000和auto_aspen_2"，auto_aspen_2的位置就改变了
    # 所以必须从后往前替换
    replacement_count = 0
    for match_start, match_end in reversed(matches):
        # 找到涉及的runs
        affected_runs = []
        for run_info in runs_info:
            if (run_info['start_pos'] < match_end and run_info['end_pos'] > match_start):
                affected_runs.append(run_info)
        
        if not affected_runs:
            continue
            
        print(f"   替换位置 {match_start}-{match_end}，涉及runs: {[r['index'] for r in affected_runs]}")
        
        # 执行跨run替换
        if len(affected_runs) == 1:
            # 单run情况
            run_info = affected_runs[0]
            run = run_info['run_obj']
            local_start = match_start - run_info['start_pos']
            local_end = match_end - run_info['start_pos']
            
            new_run_text = run.text[:local_start] + new_text + run.text[local_end:]
            run.text = new_run_text
            
            # 保持格式
            apply_formatting_to_run(run, run_info, force_font_size, old_text=old_text)
            replacement_count += 1
            
        else:
            # 多run情况：重构affected runs
            first_run = affected_runs[0]['run_obj']
            last_run = affected_runs[-1]['run_obj']
            
            # 计算每个run中的替换部分
            first_run_local_start = match_start - affected_runs[0]['start_pos']
            last_run_local_end = match_end - affected_runs[-1]['start_pos']
            
            # 构建新文本
            new_first_run_text = first_run.text[:first_run_local_start] + new_text
            new_last_run_text = last_run.text[last_run_local_end:]
            
            # 设置第一个run
            first_run.text = new_first_run_text
            apply_formatting_to_run(first_run, affected_runs[0], force_font_size, old_text=old_text)
            
            # 设置最后一个run
            last_run.text = new_last_run_text
            apply_formatting_to_run(last_run, affected_runs[-1], force_font_size, old_text=old_text)
            
            # 清空中间的runs
            for run_info in affected_runs[1:-1]:
                run_info['run_obj'].text = ""
            
            replacement_count += 1
    
    print(f"✅ 跨run替换完成: {replacement_count}次")
    return replacement_count

def apply_formatting_to_run(run, run_info, force_font_size=None, old_text=None):
    """
    应用格式到run，支持特殊变量的特殊字体大小
    
    Args:
        run: docx run对象
        run_info: run的格式信息
        force_font_size: 强制字体大小
        old_text: 被替换的原始文本，用于判断是否为特殊变量
    """
    if run_info['bold'] is not None:
        run.bold = run_info['bold']
    if run_info['italic'] is not None:
        run.italic = run_info['italic']
    if run_info['underline'] is not None:
        run.underline = run_info['underline']
    if run_info['font_name'] is not None:
        run.font.name = run_info['font_name']
    
    # 智能字体大小设置
    final_font_size = None
    font_size_source = ""
    
    # 检查是否为特殊变量
    if old_text:
        special_config = get_special_font_config()
        if old_text in special_config.get('special_font_size', {}):
            final_font_size = special_config['special_font_size'][old_text]
            font_size_source = f"特殊变量字体({final_font_size}pt)"
    
    # 如果不是特殊变量，使用常规逻辑
    if final_font_size is None:
        if force_font_size is not None:
            final_font_size = force_font_size
            font_size_source = f"强制字体({final_font_size}pt)"
        elif run_info['font_size'] is not None:
            # 处理原始字体大小
            if hasattr(run_info['font_size'], 'pt'):
                final_font_size = run_info['font_size'].pt
            else:
                final_font_size = run_info['font_size']
            font_size_source = f"原始字体({final_font_size}pt)"
    
    # 应用字体大小
    if final_font_size is not None:
        run.font.size = Pt(final_font_size)
        print(f"🔧 应用字体大小: {font_size_source}")
    
    if run_info['font_color'] is not None:
        run.font.color.rgb = run_info['font_color']

def replace_text_in_cell(cell, old_text, new_text, force_font_size=None):
    """
    在表格单元格中替换文本，保持原始格式
    
    Args:
        cell: docx表格单元格对象
        old_text (str): 要替换的文本
        new_text (str): 新文本
        force_font_size (float, optional): 强制设置字体大小（点数）
    
    Returns:
        int: 替换次数
    """
    total_replacements = 0
    for paragraph in cell.paragraphs:
        total_replacements += replace_text_in_paragraph(paragraph, old_text, new_text, force_font_size)
    return total_replacements

def replace_text_in_docx_with_formatting(docx_path, replacements, output_docx_path=None, force_font_size=None):
    """
    读取docx文件，替换指定文本并保持格式，然后保存为新的docx文件
    
    Args:
        docx_path (str): 输入的docx文件路径
        replacements (dict): 需要替换的文本字典，格式为 {old_text: new_text}
        output_docx_path (str, optional): 输出docx文件路径，如果不指定则使用原文件名加_modified后缀
        force_font_size (float, optional): 强制设置字体大小（点数），如果不提供则保持原始大小
    
    Returns:
        str: 生成的docx文件路径
    """
    print(f"正在读取文档: {docx_path}")
    if force_font_size:
        print(f"🔧 将强制设置字体大小为: {force_font_size}pt")
    
    # 加载docx文档
    doc = docx.Document(docx_path)
    
    # 按auto_aspen编号倒序排序替换键
    sorted_keys = sort_auto_aspen_keys_reverse(replacements)
    print(f"替换顺序: {sorted_keys}")
    
    # 替换段落中的文本
    replaced_count = 0
    for paragraph in doc.paragraphs:
        for old_text in sorted_keys:
            new_text = replacements[old_text]
            count = replace_text_in_paragraph(paragraph, old_text, new_text, force_font_size)
            if count > 0:
                print(f"在段落中找到并替换: '{old_text}' -> '{new_text}' ({count}次)")
                replaced_count += count
    
    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text in sorted_keys:
                    new_text = replacements[old_text]
                    count = replace_text_in_cell(cell, old_text, new_text, force_font_size)
                    if count > 0:
                        print(f"在表格中找到并替换: '{old_text}' -> '{new_text}' ({count}次)")
                        replaced_count += count
    
    print(f"总共进行了 {replaced_count} 次替换")
    
    # 确定输出docx路径
    if output_docx_path is None:
        docx_file = Path(docx_path)
        output_docx_path = docx_file.parent / f"{docx_file.stem}_modified.docx"
    
    # 保存修改后的docx文件
    doc.save(output_docx_path)
    print(f"修改后的文档已保存: {output_docx_path}")
    
    return str(output_docx_path)

# 保留原有的简单替换函数作为备选
def replace_text_in_docx(docx_path, replacements, output_docx_path=None):
    """
    读取docx文件，替换指定文本，并保存为新的docx文件（简单模式，不保持格式）
    
    Args:
        docx_path (str): 输入的docx文件路径
        replacements (dict): 需要替换的文本字典，格式为 {old_text: new_text}
        output_docx_path (str, optional): 输出docx文件路径，如果不指定则使用原文件名加_modified后缀
    
    Returns:
        str: 生成的docx文件路径
    """
    print(f"正在读取文档: {docx_path}")
    
    # 加载docx文档
    doc = docx.Document(docx_path)
    
    # 按auto_aspen编号倒序排序替换键
    sorted_keys = sort_auto_aspen_keys_reverse(replacements)
    print(f"替换顺序: {sorted_keys}")
    
    # 替换段落中的文本
    replaced_count = 0
    for paragraph in doc.paragraphs:
        for old_text in sorted_keys:
            new_text = replacements[old_text]
            if old_text in paragraph.text:
                print(f"在段落中找到并替换: '{old_text}' -> '{new_text}'")
                paragraph.text = paragraph.text.replace(old_text, new_text)
                replaced_count += 1
    
    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text in sorted_keys:
                    new_text = replacements[old_text]
                    if old_text in cell.text:
                        print(f"在表格中找到并替换: '{old_text}' -> '{new_text}'")
                        cell.text = cell.text.replace(old_text, new_text)
                        replaced_count += 1
    
    print(f"总共进行了 {replaced_count} 次替换")
    
    # 确定输出docx路径
    if output_docx_path is None:
        docx_file = Path(docx_path)
        output_docx_path = docx_file.parent / f"{docx_file.stem}_modified.docx"
    
    # 保存修改后的docx文件
    doc.save(output_docx_path)
    print(f"修改后的文档已保存: {output_docx_path}")
    
    return str(output_docx_path)

def convert_to_pdf_with_libre_office(docx_path):
    """
    尝试使用LibreOffice转换docx为PDF
    """
    output_dir = Path(docx_path).parent
    cmd = f'soffice --headless --convert-to pdf --outdir "{output_dir}" "{docx_path}"'
    print(f"尝试使用LibreOffice转换: {cmd}")
    
    try:
        result = os.system(cmd)
        if result == 0:
            pdf_path = output_dir / f"{Path(docx_path).stem}.pdf"
            if pdf_path.exists():
                print(f"成功使用LibreOffice生成PDF: {pdf_path}")
                return str(pdf_path)
        print("LibreOffice转换失败或未安装")
        return None
    except Exception as e:
        print(f"LibreOffice转换出错: {e}")
        return None

def process_document_with_parameters(docx_path, custom_parameters=None, image_replacements=None, text_to_image_replacements=None, output_docx_path=None, convert_to_pdf=True, preserve_formatting=True, force_font_size=None):
    """
    使用参数映射处理文档的主函数（支持文本替换、图片替换、文字转图片）
    
    Args:
        docx_path (str): 输入的docx文件路径
        custom_parameters (dict, optional): 自定义参数值
        image_replacements (dict, optional): 图片替换字典（图片换图片）
        text_to_image_replacements (dict, optional): 文字到图片替换字典（文字换图片）
        output_docx_path (str, optional): 输出docx文件路径
        convert_to_pdf (bool): 是否转换为PDF
        preserve_formatting (bool): 是否保持原始格式
        force_font_size (float, optional): 强制设置字体大小（点数）
    
    Returns:
        dict: 包含处理结果的字典
    """
    try:
        # 创建替换字典
        replacements = create_replacement_dict(custom_parameters)
        
        # 选择替换方法
        if preserve_formatting:
            print("使用格式保持模式进行文本替换...")
            if force_font_size:
                print(f"🔧 将强制设置字体大小为: {force_font_size}pt")
            modified_docx_path = replace_text_in_docx_with_formatting(docx_path, replacements, output_docx_path, force_font_size)
        else:
            print("使用简单模式进行文本替换...")
            modified_docx_path = replace_text_in_docx(docx_path, replacements, output_docx_path)
        
        result = {
            "success": True,
            "modified_docx_path": modified_docx_path,
            "pdf_path": None,
            "replacements_made": len(replacements),
            "images_replaced": 0,
            "text_to_image_replaced": 0
        }
        
        # 处理图片替换（图片换图片）
        if image_replacements:
            print("\n开始处理图片替换（图片换图片）...")
            image_result = replace_images_in_docx(modified_docx_path, image_replacements, modified_docx_path)
            if image_result["success"]:
                result["images_replaced"] = image_result["images_replaced"]
                print(f"成功替换 {image_result['images_replaced']} 个图片")
            else:
                print(f"图片替换失败: {image_result['error']}")
        
        # 处理文字到图片替换
        if text_to_image_replacements:
            print("\n开始处理文字到图片替换...")
            text_to_image_result = replace_text_with_images_in_docx(modified_docx_path, text_to_image_replacements, modified_docx_path)
            if text_to_image_result["success"]:
                result["text_to_image_replaced"] = text_to_image_result["text_to_image_replacements"]
                print(f"成功将 {text_to_image_result['text_to_image_replacements']} 处文字替换为图片")
            else:
                print(f"文字到图片替换失败: {text_to_image_result['error']}")
        
        # 尝试转换为PDF
        if convert_to_pdf:
            print("\n尝试转换为PDF...")
            pdf_path = convert_to_pdf_with_libre_office(modified_docx_path)
            result["pdf_path"] = pdf_path
        
        return result
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "modified_docx_path": None,
            "pdf_path": None,
            "replacements_made": 0,
            "images_replaced": 0,
            "text_to_image_replaced": 0
        }

def main():
    """主函数：演示如何使用参数映射功能"""
    docx_path = "models/RE_template.docx"  # 使用模板文件
    
    if not os.path.exists(docx_path):
        print(f"错误：找不到文件 {docx_path}")
        return
    
    # 自定义一些参数值
    custom_parameters = {
        "auto_aspen_1": "45000",    # 最大气量
        "auto_aspen_2": "13.8",     # 进站压力
        "auto_aspen_3": "30",       # 平均进气温度
        "auto_aspen_5": "680",      # 净发电功率
    }
    
    print("开始处理文档...")
    result = process_document_with_parameters(
        docx_path, 
        custom_parameters,
        convert_to_pdf=True
    )
    
    if result["success"]:
        print(f"\n文档处理成功！")
        print(f"修改后的docx文件: {result['modified_docx_path']}")
        if result["pdf_path"]:
            print(f"PDF文件: {result['pdf_path']}")
        print(f"共进行了 {result['replacements_made']} 个参数的替换")
    else:
        print(f"\n文档处理失败: {result['error']}")

def generate_document(parameters=None, images=None, text_to_images=None, output_name=None, convert_pdf=True, preserve_formatting=True, force_font_size=None):
    """
    简化的API函数：生成带有指定参数、图片替换、文字转图片的文档
    
    Args:
        parameters (dict, optional): 自定义参数值，格式为 {"auto_aspen_1": "value", ...}
        images (dict, optional): 图片替换字典（图片换图片），格式为 {"old_image.png": {"new_path": "path", "width": w, "height": h}}
        text_to_images (dict, optional): 文字到图片替换字典，格式为 {"文字": {"image_path": "path", "width": w, "height": h}}
        output_name (str, optional): 输出文件名（不含扩展名），默认为"RE_generated"
        convert_pdf (bool): 是否同时生成PDF文件
        preserve_formatting (bool): 是否保持原始格式（字体、颜色等）
        force_font_size (float, optional): 强制设置字体大小（点数），如12.0表示12pt字体
    
    Returns:
        dict: 包含生成文件路径的结果字典
    
    Example:
        # 使用默认参数生成文档，保持格式
        result = generate_document()
        
        # 使用自定义参数生成文档，保持格式
        params = {
            "auto_aspen_1": "60000",  # 最大气量
            "auto_aspen_2": "14.0",   # 进站压力
            "auto_aspen_5": "800"     # 净发电功率
        }
        result = generate_document(parameters=params, output_name="custom_report", preserve_formatting=True)
        
        # 强制设置字体大小为12pt
        result = generate_document(parameters=params, force_font_size=12.0)
        
        # 同时替换参数和图片
        images = {
            "chart1.png": {
                "new_path": "static/re/new_chart.png",
                "width": 6.0,
                "height": 4.0
            }
        }
        result = generate_document(parameters=params, images=images, output_name="full_custom")
        
        # 将文字替换为图片
        text_to_images = {
            "[图表1]": {
                "image_path": "models/demo.png",
                "width": 5.0,
                "height": 3.0
            },
            "【插入图片】": {
                "image_path": "models/demo.png",
                "width": 4.0,
                "height": 2.5
            }
        }
        result = generate_document(parameters=params, text_to_images=text_to_images, output_name="text_to_image")
        
        # 使用简单模式（不保持格式，但速度更快）
        result = generate_document(parameters=params, preserve_formatting=False)
    """
    template_path = "models/RE_template.docx"
    
    if not os.path.exists(template_path):
        return {
            "success": False,
            "error": f"模板文件不存在: {template_path}",
            "docx_path": None,
            "pdf_path": None
        }
    
    # 确定输出文件名
    if output_name is None:
        output_name = "RE_generated"
    
    output_docx_path = f"{file_dir}/{output_name}.docx"
    
    # 处理文档
    result = process_document_with_parameters(
        template_path,
        parameters,
        images,
        text_to_images,
        output_docx_path,
        convert_pdf,
        preserve_formatting,
        force_font_size
    )
    
    return {
        "success": result["success"],
        "error": result.get("error"),
        "docx_path": result["modified_docx_path"],
        "pdf_path": result["pdf_path"],
        "parameters_replaced": result["replacements_made"],
        "images_replaced": result["images_replaced"],
        "text_to_image_replaced": result["text_to_image_replaced"]
    }

def sort_auto_aspen_keys_reverse(replacements):
    """
    对替换键进行智能排序，适应跨run替换需求
    
    排序规则：
    1. 首先按key长度倒序排序（长key优先，避免短key污染长key）
    2. 长度相同时按字典序倒序排序（确保后出现的key先替换）
    3. 非auto_aspen参数保持原顺序
    
    这样的排序策略确保了：
    - auto_aspen_623 比 auto_aspen_6 先替换（避免污染）
    - auto_aspen_time 比 auto_aspen_26 先替换（长度优先）
    - auto_aspen_9 比 auto_aspen_8 先替换（倒序，适应跨run替换）
    
    Args:
        replacements (dict): 替换字典
    
    Returns:
        list: 按智能规则排序的键列表
    """
    auto_aspen_keys = []
    other_keys = []
    
    for key in replacements.keys():
        if key.startswith("auto_aspen_"):
            auto_aspen_keys.append(key)
        else:
            other_keys.append(key)
    
    # 智能排序策略：
    # 1. 首先按长度倒序（长key优先）
    # 2. 长度相同时按字典序倒序（后面的key先替换，适应跨run替换）
    auto_aspen_keys.sort(key=lambda x: (-len(x), x), reverse=True)

    # 默认不刷屏；需要调试占位符替换时设置 AUTO_ASPEN_DEBUG_REPLACEMENTS=1
    if os.getenv("AUTO_ASPEN_DEBUG_REPLACEMENTS", "").strip() in ("1", "true", "yes"):
        print(f"🔍 跨run适配排序结果: {auto_aspen_keys[:10]}...")
        if len(auto_aspen_keys) > 0:
            print("🧪 排序逻辑验证: 长度优先 → 字典序倒序")
            length_groups = {}
            for key in auto_aspen_keys[:8]:
                length_groups.setdefault(len(key), []).append(key)
            for length in sorted(length_groups.keys(), reverse=True):
                print(f"   长度{length}: {length_groups[length]}")

    # 返回排序后的完整键列表
    return auto_aspen_keys + other_keys

def find_images_in_document(doc):
    """
    查找文档中的所有图片
    
    Args:
        doc: docx文档对象
    
    Returns:
        list: 包含图片信息的列表，每个元素为 {'paragraph': paragraph, 'run': run, 'inline_shape': shape, 'image_name': name}
    """
    images = []
    
    # 遍历段落中的图片
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for inline_shape in run.element.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                # 获取图片的关系ID
                rId = inline_shape.get(qn('r:embed'))
                if rId:
                    try:
                        # 获取图片文件名
                        image_part = doc.part.related_parts[rId]
                        image_name = os.path.basename(image_part.partname)
                        
                        images.append({
                            'paragraph': paragraph,
                            'run': run,
                            'rId': rId,
                            'image_name': image_name,
                            'image_part': image_part
                        })
                    except:
                        continue
    
    # 遍历表格中的图片
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for inline_shape in run.element.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            rId = inline_shape.get(qn('r:embed'))
                            if rId:
                                try:
                                    image_part = doc.part.related_parts[rId]
                                    image_name = os.path.basename(image_part.partname)
                                    
                                    images.append({
                                        'paragraph': paragraph,
                                        'run': run,
                                        'rId': rId,
                                        'image_name': image_name,
                                        'image_part': image_part
                                    })
                                except:
                                    continue
    
    return images

def replace_image_by_name(doc, old_image_name, new_image_path, width=None, height=None):
    """
    根据图片名称替换文档中的图片
    
    Args:
        doc: docx文档对象
        old_image_name (str): 要替换的图片名称（如 "image1.png"）
        new_image_path (str): 新图片的文件路径
        width (float, optional): 新图片宽度（英寸）
        height (float, optional): 新图片高度（英寸）
    
    Returns:
        int: 替换的图片数量
    """
    if not os.path.exists(new_image_path):
        print(f"错误：新图片文件不存在 {new_image_path}")
        return 0
    
    replaced_count = 0
    images = find_images_in_document(doc)
    
    for image_info in images:
        if old_image_name in image_info['image_name'] or image_info['image_name'].endswith(old_image_name):
            try:
                # 删除原有图片
                run = image_info['run']
                
                # 清除run中的所有内容
                run.clear()
                
                # 添加新图片
                if width and height:
                    run.add_picture(new_image_path, width=Inches(width), height=Inches(height))
                else:
                    run.add_picture(new_image_path)
                
                print(f"成功替换图片: {image_info['image_name']} -> {new_image_path}")
                replaced_count += 1
                
            except Exception as e:
                print(f"替换图片时出错: {e}")
    
    return replaced_count

def replace_image_by_position(doc, paragraph_index, new_image_path, width=None, height=None):
    """
    根据段落位置替换图片
    
    Args:
        doc: docx文档对象
        paragraph_index (int): 段落索引（从0开始）
        new_image_path (str): 新图片的文件路径
        width (float, optional): 新图片宽度（英寸）
        height (float, optional): 新图片高度（英寸）
    
    Returns:
        bool: 是否成功替换
    """
    if not os.path.exists(new_image_path):
        print(f"错误：新图片文件不存在 {new_image_path}")
        return False
    
    try:
        if paragraph_index >= len(doc.paragraphs):
            print(f"错误：段落索引超出范围 {paragraph_index}")
            return False
        
        paragraph = doc.paragraphs[paragraph_index]
        
        # 查找该段落中的图片
        image_found = False
        for run in paragraph.runs:
            if run.element.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                # 清除run中的所有内容
                run.clear()
                
                # 添加新图片
                if width and height:
                    run.add_picture(new_image_path, width=Inches(width), height=Inches(height))
                else:
                    run.add_picture(new_image_path)
                
                print(f"成功在段落 {paragraph_index} 位置替换图片: {new_image_path}")
                image_found = True
                break
        
        if not image_found:
            print(f"在段落 {paragraph_index} 中未找到图片")
            return False
        
        return True
        
    except Exception as e:
        print(f"替换图片时出错: {e}")
        return False

def add_image_to_paragraph(doc, paragraph_index, image_path, width=None, height=None):
    """
    在指定段落添加图片
    
    Args:
        doc: docx文档对象
        paragraph_index (int): 段落索引（从0开始）
        image_path (str): 图片文件路径
        width (float, optional): 图片宽度（英寸）
        height (float, optional): 图片高度（英寸）
    
    Returns:
        bool: 是否成功添加
    """
    if not os.path.exists(image_path):
        print(f"错误：图片文件不存在 {image_path}")
        return False
    
    try:
        if paragraph_index >= len(doc.paragraphs):
            print(f"错误：段落索引超出范围 {paragraph_index}")
            return False
        
        paragraph = doc.paragraphs[paragraph_index]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        
        # 添加图片
        if width and height:
            run.add_picture(image_path, width=Inches(width), height=Inches(height))
        else:
            run.add_picture(image_path)
        
        print(f"成功在段落 {paragraph_index} 添加图片: {image_path}")
        return True
        
    except Exception as e:
        print(f"添加图片时出错: {e}")
        return False

def create_image_replacement_dict():
    """
    创建图片替换映射字典的示例
    
    Returns:
        dict: 图片替换映射字典
    """
    return {
        # 格式: "原图片名称": {"new_path": "新图片路径", "width": 宽度, "height": 高度}
        "chart1.png": {
            "new_path": "static/re/new_chart1.png",
            "width": 6.0,  # 英寸
            "height": 4.0
        },
        "diagram1.jpg": {
            "new_path": "static/re/new_diagram1.jpg",
            "width": 5.0,
            "height": 3.5
        }
    }

def replace_images_in_docx(docx_path, image_replacements, output_docx_path=None):
    """
    批量替换docx文档中的图片
    
    Args:
        docx_path (str): 输入的docx文件路径
        image_replacements (dict): 图片替换字典，格式为 {image_name: {"new_path": path, "width": w, "height": h}}
        output_docx_path (str, optional): 输出docx文件路径
    
    Returns:
        dict: 处理结果
    """
    print(f"正在读取文档进行图片替换: {docx_path}")
    
    try:
        # 加载docx文档
        doc = docx.Document(docx_path)
        
        # 查找所有图片
        images = find_images_in_document(doc)
        print(f"文档中找到 {len(images)} 个图片")
        
        replaced_count = 0
        
        # 执行图片替换
        for old_name, replacement_info in image_replacements.items():
            new_path = replacement_info.get("new_path")
            width = replacement_info.get("width")
            height = replacement_info.get("height")
            
            if not new_path or not os.path.exists(new_path):
                print(f"跳过 {old_name}: 新图片路径无效或文件不存在")
                continue
            
            count = replace_image_by_name(doc, old_name, new_path, width, height)
            replaced_count += count
        
        # 确定输出路径
        if output_docx_path is None:
            docx_file = Path(docx_path)
            output_docx_path = docx_file.parent / f"{docx_file.stem}_images_replaced.docx"
        
        # 保存修改后的文档
        doc.save(output_docx_path)
        print(f"图片替换完成，文档已保存: {output_docx_path}")
        
        return {
            "success": True,
            "modified_docx_path": str(output_docx_path),
            "images_replaced": replaced_count,
            "total_images_found": len(images)
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "images_replaced": 0
        }

def replace_text_with_image(paragraph, old_text, image_path, width=None, height=None):
    """
    在段落中将文字替换为图片
    
    Args:
        paragraph: docx段落对象
        old_text (str): 要替换的文字
        image_path (str): 图片路径
        width (float, optional): 图片宽度（英寸）
        height (float, optional): 图片高度（英寸）
    
    Returns:
        int: 替换次数
    """
    if old_text not in paragraph.text:
        return 0
    
    if not os.path.exists(image_path):
        print(f"错误：图片文件不存在 {image_path}")
        return 0
    
    try:
        # 获取段落的完整文本
        full_text = paragraph.text
        
        if old_text not in full_text:
            return 0
        
        # 计算替换次数
        replacement_count = full_text.count(old_text)
        
        # 清空段落中的所有runs
        for run in paragraph.runs:
            run.clear()
        
        # 按old_text分割文本
        parts = full_text.split(old_text)
        
        # 重新构建段落内容
        for i, part in enumerate(parts):
            if i > 0:  # 在每个分割点插入图片
                # 添加图片
                run = paragraph.add_run()
                if width and height:
                    run.add_picture(image_path, width=Inches(width), height=Inches(height))
                else:
                    run.add_picture(image_path)
            
            if part:  # 添加文本部分
                paragraph.add_run(part)
        
        print(f"成功将文字 '{old_text}' 替换为图片 {image_path} ({replacement_count}次)")
        return replacement_count
        
    except Exception as e:
        print(f"将文字替换为图片时出错: {e}")
        return 0

def replace_text_with_image_in_cell(cell, old_text, image_path, width=None, height=None):
    """
    在表格单元格中将文字替换为图片
    
    Args:
        cell: docx表格单元格对象
        old_text (str): 要替换的文字
        image_path (str): 图片路径
        width (float, optional): 图片宽度（英寸）
        height (float, optional): 图片高度（英寸）
    
    Returns:
        int: 替换次数
    """
    total_replacements = 0
    for paragraph in cell.paragraphs:
        total_replacements += replace_text_with_image(paragraph, old_text, image_path, width, height)
    return total_replacements

def replace_text_with_images_in_docx(docx_path, text_to_image_replacements, output_docx_path=None):
    """
    将docx文档中的指定文字替换为图片
    
    Args:
        docx_path (str): 输入的docx文件路径
        text_to_image_replacements (dict): 文字到图片的替换字典
            格式: {text: {"image_path": path, "width": w, "height": h}}
        output_docx_path (str, optional): 输出docx文件路径
    
    Returns:
        dict: 处理结果
    """
    print(f"正在读取文档进行文字到图片替换: {docx_path}")
    
    try:
        # 加载docx文档
        doc = docx.Document(docx_path)
        
        total_replacements = 0
        
        # 先按文字长度倒序排序，避免短文字被长文字影响
        sorted_texts = sorted(text_to_image_replacements.keys(), key=len, reverse=True)
        
        # 替换段落中的文字
        for paragraph in doc.paragraphs:
            for text in sorted_texts:
                replacement_info = text_to_image_replacements[text]
                image_path = replacement_info.get("image_path")
                width = replacement_info.get("width")
                height = replacement_info.get("height")
                
                if not image_path:
                    continue
                
                count = replace_text_with_image(paragraph, text, image_path, width, height)
                total_replacements += count
        
        # 替换表格中的文字
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for text in sorted_texts:
                        replacement_info = text_to_image_replacements[text]
                        image_path = replacement_info.get("image_path")
                        width = replacement_info.get("width")
                        height = replacement_info.get("height")
                        
                        if not image_path:
                            continue
                        
                        count = replace_text_with_image_in_cell(cell, text, image_path, width, height)
                        total_replacements += count
        
        # 确定输出路径
        if output_docx_path is None:
            docx_file = Path(docx_path)
            output_docx_path = docx_file.parent / f"{docx_file.stem}_text_to_images.docx"
        
        # 保存修改后的文档
        doc.save(output_docx_path)
        print(f"文字到图片替换完成，共替换 {total_replacements} 次，文档已保存: {output_docx_path}")
        
        return {
            "success": True,
            "modified_docx_path": str(output_docx_path),
            "text_to_image_replacements": total_replacements
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "text_to_image_replacements": 0
        }

def create_text_to_image_mapping():
    """
    创建文字到图片替换映射的示例
    
    Returns:
        dict: 文字到图片的替换字典
    """
    return {
        # 格式: "要替换的文字": {"image_path": "图片路径", "width": 宽度, "height": 高度}
        "[图表1]": {
            "image_path": "models/demo.png",
            "width": 5.0,
            "height": 3.0
        },
        "[机组示意图]": {
            "image_path": "models/demo.png",
            "width": 6.0,
            "height": 4.0
        },
        "[示意图]": {
            "image_path": "models/demo.png",
            "width": 4.5,
            "height": 3.5
        },
        "【插入图片】": {
            "image_path": "models/demo.png",
            "width": 3.0,
            "height": 2.0
        }
    }

if __name__ == "__main__":
    # 运行测试
    
    print("\n" + "="*50)
    print("测试文本替换功能")
    print("="*50 + "\n")
    
    custom_params = get_auto_aspen_parameter_mapping()
    result1 = generate_document(parameters=custom_params, output_name="demo_custom_formatted", convert_pdf=False, 
    preserve_formatting=True)
    print(f"文本替换结果: {result1}")
    
    # print("\n" + "="*50)
    # print("测试图片替换功能")
    # print("="*50 + "\n")
    
    # 检查模板文件中的图片
    # template_path = "models/RE_template.docx"
    
    # # 使用demo.png作为替换图片
    # demo_image = "models/demo.png"
    # if os.path.exists(demo_image):
    #     text_to_image_replacements = {
    #         "auto_aspen_image_1": {
    #             "image_path": "models/demo.png",
    #             "width": 5.0,
    #             "height": 3.0
    #             },
    #     }
    #     # 测试全功能（文本+图片+文字转图片）
    #     print(f"\n全功能测试（文本+图片+文字转图片）:")
    #     full_result = generate_document(
    #         parameters=custom_params,
    #         images=[],
    #         text_to_images=text_to_image_replacements,
    #         output_name="demo_full_features",
    #         convert_pdf=False,
    #         preserve_formatting=False
    #     )
    #     print(f"全功能测试结果: {full_result}")
        
    # else:
    #     print(f"\n警告: Demo图片不存在: {demo_image}")
    #     print("将跳过图片相关测试")
    
    # print("\n" + "="*50)
    # print("测试完成")
    # print("="*50)
    
